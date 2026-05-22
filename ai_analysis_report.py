import json
import math
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from xml.sax.saxutils import escape as xml_escape

import requests


AI_PROVIDER_DISABLED = "Disabled"
AI_PROVIDER_OPENAI = "ChatGPT"
AI_PROVIDER_ANTHROPIC = "Claude"
AI_PROVIDER_VALUES = (AI_PROVIDER_DISABLED, AI_PROVIDER_OPENAI, AI_PROVIDER_ANTHROPIC)
AI_ANALYSIS_TIMEOUT_SECONDS = 120
AI_ANALYSIS_MAX_TRANSACTIONS = 250
AI_ANALYSIS_MAX_IO = 20
AI_ANALYSIS_PROMPT = (
    "Analyze the data to detect patterns consistent with money laundering behavior or illicit finance behavior. "
    "Output a trend analysis showing interractive graphs of transaction volume and net worth of the account over "
    "time via plotly. Output an interractive sankey graph via plotly for showing where the money comes from and "
    "goes to in relation to the specified wallet."
)
DEFAULT_OPENAI_MODEL = "gpt-5-mini"
DEFAULT_ANTHROPIC_MODEL = "claude-sonnet-4-20250514"


class AIAnalysisError(RuntimeError):
    pass


def generate_ai_analysis_report(
    result: Dict[str, Any],
    seed: str,
    mode: str,
    provider: str,
    api_token: str,
    output_dir: Path,
) -> Dict[str, Path]:
    provider = str(provider or "").strip()
    token = str(api_token or "").strip()
    if provider not in {AI_PROVIDER_OPENAI, AI_PROVIDER_ANTHROPIC}:
        raise AIAnalysisError("Select ChatGPT or Claude to run AI analysis.")
    if not token:
        raise AIAnalysisError(f"{provider} API token is required for AI analysis.")

    output_dir.mkdir(parents=True, exist_ok=True)
    focus_wallet = _resolve_focus_wallet(result=result, seed=seed, mode=mode)
    series = _build_time_series(result=result, focus_wallet=focus_wallet)
    sankey_links = _build_sankey_links(result=result, focus_wallet=focus_wallet)
    chart_paths = _write_plotly_charts(output_dir=output_dir, series=series, sankey_links=sankey_links)
    compact_payload = _compact_trace_for_ai(
        result=result,
        seed=seed,
        mode=mode,
        focus_wallet=focus_wallet,
        series=series,
        sankey_links=sankey_links,
    )

    prompt = _build_ai_prompt(compact_payload=compact_payload, chart_paths=chart_paths)
    ai_text = _call_ai_provider(provider=provider, api_token=token, prompt=prompt)

    text_path = output_dir / "ai_analysis.txt"
    text_path.write_text(ai_text, encoding="utf-8")
    docx_path = _write_docx_report(
        output_path=output_dir / "ai_analysis.docx",
        provider=provider,
        focus_wallet=focus_wallet,
        ai_text=ai_text,
        chart_paths=chart_paths,
        compact_payload=compact_payload,
    )

    return {
        "docx": docx_path,
        "text": text_path,
        "trend_chart": chart_paths["trend_chart"],
        "sankey_chart": chart_paths["sankey_chart"],
    }


def _build_ai_prompt(compact_payload: Dict[str, Any], chart_paths: Dict[str, Path]) -> str:
    chart_manifest = {
        name: str(path.name)
        for name, path in chart_paths.items()
    }
    return (
        f"{AI_ANALYSIS_PROMPT}\n\n"
        "Use the supplied transaction trace summary and chart manifest. The application will generate the "
        "interactive Plotly HTML files listed in the manifest, so your response should provide the written "
        "trend/risk analysis, call out what the graphs should help the analyst inspect, and avoid making a "
        "legal conclusion. Label hypotheses as indicators requiring analyst validation.\n\n"
        f"Plotly chart manifest:\n{json.dumps(chart_manifest, indent=2)}\n\n"
        f"Transaction trace summary:\n{json.dumps(compact_payload, indent=2)}"
    )


def _call_ai_provider(provider: str, api_token: str, prompt: str) -> str:
    if provider == AI_PROVIDER_OPENAI:
        return _call_openai(api_token=api_token, prompt=prompt)
    if provider == AI_PROVIDER_ANTHROPIC:
        return _call_anthropic(api_token=api_token, prompt=prompt)
    raise AIAnalysisError(f"Unsupported AI provider: {provider}")


def _call_openai(api_token: str, prompt: str) -> str:
    response = requests.post(
        "https://api.openai.com/v1/responses",
        headers={
            "Authorization": f"Bearer {api_token}",
            "Content-Type": "application/json",
        },
        json={
            "model": DEFAULT_OPENAI_MODEL,
            "input": prompt,
            "max_output_tokens": 4000,
        },
        timeout=AI_ANALYSIS_TIMEOUT_SECONDS,
    )
    if response.status_code >= 400:
        raise AIAnalysisError(f"ChatGPT API error HTTP {response.status_code}: {response.text[:500]}")
    payload = response.json()
    text = _extract_openai_text(payload)
    if not text:
        raise AIAnalysisError("ChatGPT API returned no text output.")
    return text


def _extract_openai_text(payload: Dict[str, Any]) -> str:
    direct = payload.get("output_text")
    if isinstance(direct, str) and direct.strip():
        return direct.strip()

    parts: List[str] = []
    for output in payload.get("output", []) if isinstance(payload.get("output"), list) else []:
        if not isinstance(output, dict):
            continue
        for item in output.get("content", []) if isinstance(output.get("content"), list) else []:
            if not isinstance(item, dict):
                continue
            text = item.get("text")
            if isinstance(text, str) and text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts).strip()


def _call_anthropic(api_token: str, prompt: str) -> str:
    response = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_token,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
        },
        json={
            "model": DEFAULT_ANTHROPIC_MODEL,
            "max_tokens": 4000,
            "messages": [{"role": "user", "content": prompt}],
        },
        timeout=AI_ANALYSIS_TIMEOUT_SECONDS,
    )
    if response.status_code >= 400:
        raise AIAnalysisError(f"Claude API error HTTP {response.status_code}: {response.text[:500]}")
    payload = response.json()
    parts = [
        item.get("text", "").strip()
        for item in payload.get("content", [])
        if isinstance(item, dict) and item.get("type") == "text" and str(item.get("text", "")).strip()
    ]
    text = "\n\n".join(parts).strip()
    if not text:
        raise AIAnalysisError("Claude API returned no text output.")
    return text


def _resolve_focus_wallet(result: Dict[str, Any], seed: str, mode: str) -> str:
    if mode == "address" and seed:
        return str(seed).strip()

    addresses: Counter = Counter()
    for tx in result.get("transactions", []):
        for item in tx.get("inputs", []) + tx.get("outputs", []):
            address = str(item.get("address", "")).strip()
            if address and address.lower() != "unknown":
                addresses[address] += 1
    return addresses.most_common(1)[0][0] if addresses else str(seed or "").strip()


def _build_time_series(result: Dict[str, Any], focus_wallet: str) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    focus_norm = _normalize_address(focus_wallet)
    running_balance = 0.0

    tx_rows = []
    for tx in result.get("transactions", []):
        timestamp = _parse_timestamp(tx.get("timestamp"))
        inputs = tx.get("inputs", [])
        outputs = tx.get("outputs", [])
        input_total = _sum_amounts(inputs)
        output_total = _sum_amounts(outputs)
        inbound = _sum_for_address(outputs, focus_norm)
        outbound = _sum_for_address(inputs, focus_norm)
        if focus_norm and inbound == 0 and outbound == 0:
            net_flow = output_total - input_total
        else:
            net_flow = inbound - outbound
        volume = max(input_total, output_total, inbound + outbound)
        tx_rows.append(
            {
                "timestamp": timestamp,
                "txid": str(tx.get("txid", "")),
                "volume": _round_float(volume),
                "inbound": _round_float(inbound),
                "outbound": _round_float(outbound),
                "net_flow": _round_float(net_flow),
            }
        )

    tx_rows.sort(key=lambda item: item["timestamp"])
    for row in tx_rows:
        running_balance += row["net_flow"]
        rows.append(
            {
                "timestamp": row["timestamp"].isoformat(),
                "txid": row["txid"],
                "volume": row["volume"],
                "inbound": row["inbound"],
                "outbound": row["outbound"],
                "net_flow": row["net_flow"],
                "estimated_net_worth": _round_float(running_balance),
            }
        )
    return rows


def _build_sankey_links(result: Dict[str, Any], focus_wallet: str) -> List[Dict[str, Any]]:
    focus_norm = _normalize_address(focus_wallet)
    links: Dict[Tuple[str, str], float] = defaultdict(float)

    for tx in result.get("transactions", []):
        inputs = tx.get("inputs", [])
        outputs = tx.get("outputs", [])
        input_focus = _items_for_address(inputs, focus_norm)
        output_focus = _items_for_address(outputs, focus_norm)

        if focus_norm and output_focus:
            for txin in inputs:
                source = _display_address(txin.get("address"))
                amount = float(txin.get("amount", 0) or 0)
                if source != _display_address(focus_wallet) and amount > 0:
                    links[(source, "Specified wallet")] += amount
        if focus_norm and input_focus:
            for txout in outputs:
                target = _display_address(txout.get("address"))
                amount = float(txout.get("amount", 0) or 0)
                if target != _display_address(focus_wallet) and amount > 0:
                    links[("Specified wallet", target)] += amount
        if not focus_norm:
            for txin in inputs[:AI_ANALYSIS_MAX_IO]:
                for txout in outputs[:AI_ANALYSIS_MAX_IO]:
                    amount = min(float(txin.get("amount", 0) or 0), float(txout.get("amount", 0) or 0))
                    if amount > 0:
                        links[(_display_address(txin.get("address")), _display_address(txout.get("address")))] += amount

    top_links = sorted(links.items(), key=lambda item: item[1], reverse=True)[:60]
    return [
        {"source": source, "target": target, "amount": _round_float(amount)}
        for (source, target), amount in top_links
        if amount > 0
    ]


def _write_plotly_charts(
    output_dir: Path,
    series: List[Dict[str, Any]],
    sankey_links: List[Dict[str, Any]],
) -> Dict[str, Path]:
    try:
        import plotly.graph_objects as go
        from plotly.subplots import make_subplots
    except Exception as exc:
        return _write_plotly_cdn_charts(output_dir=output_dir, series=series, sankey_links=sankey_links)

    trend_path = output_dir / "ai_trend_analysis.html"
    sankey_path = output_dir / "ai_money_flow_sankey.html"

    x_values = [row["timestamp"] for row in series]
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    fig.add_trace(
        go.Bar(x=x_values, y=[row["volume"] for row in series], name="Transaction volume"),
        secondary_y=False,
    )
    fig.add_trace(
        go.Scatter(
            x=x_values,
            y=[row["estimated_net_worth"] for row in series],
            name="Estimated net worth",
            mode="lines+markers",
        ),
        secondary_y=True,
    )
    fig.update_layout(
        title="Transaction Volume and Estimated Net Worth Over Time",
        hovermode="x unified",
        template="plotly_white",
    )
    fig.update_xaxes(title_text="Time")
    fig.update_yaxes(title_text="Transaction volume", secondary_y=False)
    fig.update_yaxes(title_text="Estimated net worth", secondary_y=True)
    fig.write_html(str(trend_path), include_plotlyjs=True, full_html=True)

    labels = sorted({link["source"] for link in sankey_links} | {link["target"] for link in sankey_links})
    label_index = {label: idx for idx, label in enumerate(labels)}
    sankey_fig = go.Figure(
        data=[
            go.Sankey(
                node={"label": labels, "pad": 15, "thickness": 18},
                link={
                    "source": [label_index[link["source"]] for link in sankey_links],
                    "target": [label_index[link["target"]] for link in sankey_links],
                    "value": [link["amount"] for link in sankey_links],
                },
            )
        ]
    )
    sankey_fig.update_layout(title="Money Flow Sankey Relative to Specified Wallet", template="plotly_white")
    sankey_fig.write_html(str(sankey_path), include_plotlyjs=True, full_html=True)

    return {"trend_chart": trend_path, "sankey_chart": sankey_path}


def _write_plotly_cdn_charts(
    output_dir: Path,
    series: List[Dict[str, Any]],
    sankey_links: List[Dict[str, Any]],
) -> Dict[str, Path]:
    trend_path = output_dir / "ai_trend_analysis.html"
    sankey_path = output_dir / "ai_money_flow_sankey.html"

    x_values = [row["timestamp"] for row in series]
    trend_data = [
        {
            "type": "bar",
            "x": x_values,
            "y": [row["volume"] for row in series],
            "name": "Transaction volume",
        },
        {
            "type": "scatter",
            "x": x_values,
            "y": [row["estimated_net_worth"] for row in series],
            "name": "Estimated net worth",
            "mode": "lines+markers",
            "yaxis": "y2",
        },
    ]
    trend_layout = {
        "title": "Transaction Volume and Estimated Net Worth Over Time",
        "hovermode": "x unified",
        "xaxis": {"title": "Time"},
        "yaxis": {"title": "Transaction volume"},
        "yaxis2": {"title": "Estimated net worth", "overlaying": "y", "side": "right"},
        "template": "plotly_white",
    }
    _write_plotly_html(trend_path, "AI Trend Analysis", trend_data, trend_layout)

    labels = sorted({link["source"] for link in sankey_links} | {link["target"] for link in sankey_links})
    if not labels:
        labels = ["No linked flows detected"]
    label_index = {label: idx for idx, label in enumerate(labels)}
    sankey_data = [
        {
            "type": "sankey",
            "node": {"label": labels, "pad": 15, "thickness": 18},
            "link": {
                "source": [label_index[link["source"]] for link in sankey_links],
                "target": [label_index[link["target"]] for link in sankey_links],
                "value": [link["amount"] for link in sankey_links],
            },
        }
    ]
    sankey_layout = {
        "title": "Money Flow Sankey Relative to Specified Wallet",
        "template": "plotly_white",
    }
    _write_plotly_html(sankey_path, "AI Money Flow Sankey", sankey_data, sankey_layout)

    return {"trend_chart": trend_path, "sankey_chart": sankey_path}


def _write_plotly_html(path: Path, title: str, data: List[Dict[str, Any]], layout: Dict[str, Any]) -> None:
    html = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{xml_escape(title)}</title>
  <script src="https://cdn.plot.ly/plotly-2.35.2.min.js"></script>
  <style>
    body {{ margin: 0; font-family: Arial, sans-serif; }}
    #chart {{ width: 100vw; height: 100vh; }}
  </style>
</head>
<body>
  <div id="chart"></div>
  <script>
    Plotly.newPlot("chart", {json.dumps(data)}, {json.dumps(layout)}, {{responsive: true}});
  </script>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


def _compact_trace_for_ai(
    result: Dict[str, Any],
    seed: str,
    mode: str,
    focus_wallet: str,
    series: List[Dict[str, Any]],
    sankey_links: List[Dict[str, Any]],
) -> Dict[str, Any]:
    transactions = result.get("transactions", [])
    compact_transactions = []
    for tx in transactions[:AI_ANALYSIS_MAX_TRANSACTIONS]:
        compact_transactions.append(
            {
                "txid": tx.get("txid"),
                "timestamp": tx.get("timestamp"),
                "fee": tx.get("fee"),
                "input_total": _round_float(_sum_amounts(tx.get("inputs", []))),
                "output_total": _round_float(_sum_amounts(tx.get("outputs", []))),
                "inputs": _compact_io(tx.get("inputs", [])),
                "outputs": _compact_io(tx.get("outputs", [])),
            }
        )

    return {
        "metadata": result.get("metadata", {}),
        "seed": seed,
        "mode": mode,
        "focus_wallet": focus_wallet,
        "transaction_count_supplied": len(transactions),
        "transaction_count_in_ai_payload": len(compact_transactions),
        "payload_truncated": len(transactions) > len(compact_transactions),
        "time_series": series,
        "top_sankey_links": sankey_links[:60],
        "findings": result.get("findings", [])[:100],
        "transactions": compact_transactions,
    }


def _compact_io(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [
        {"address": item.get("address"), "amount": item.get("amount", 0)}
        for item in items[:AI_ANALYSIS_MAX_IO]
    ]


def _write_docx_report(
    output_path: Path,
    provider: str,
    focus_wallet: str,
    ai_text: str,
    chart_paths: Dict[str, Path],
    compact_payload: Dict[str, Any],
) -> Path:
    try:
        from docx import Document
    except Exception as exc:
        return _write_minimal_docx_report(
            output_path=output_path,
            provider=provider,
            focus_wallet=focus_wallet,
            ai_text=ai_text,
            chart_paths=chart_paths,
            compact_payload=compact_payload,
        )

    doc = Document()
    doc.add_heading("AI Transaction Analysis", 0)
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
    doc.add_paragraph(f"AI provider: {provider}")
    doc.add_paragraph(f"Specified wallet / focus: {focus_wallet or 'Not available'}")

    doc.add_heading("Interactive Plotly Graphs", level=1)
    doc.add_paragraph("Open these HTML files from the export folder to inspect the interactive charts.")
    _add_hyperlink_paragraph(doc, "Transaction volume and estimated net worth", chart_paths["trend_chart"])
    _add_hyperlink_paragraph(doc, "Money flow Sankey", chart_paths["sankey_chart"])

    doc.add_heading("AI Analysis", level=1)
    for block in ai_text.split("\n\n"):
        text = block.strip()
        if text:
            doc.add_paragraph(text)

    doc.add_heading("Trace Scope", level=1)
    doc.add_paragraph(f"Transactions supplied to app: {compact_payload.get('transaction_count_supplied')}")
    doc.add_paragraph(f"Transactions included in AI payload: {compact_payload.get('transaction_count_in_ai_payload')}")
    if compact_payload.get("payload_truncated"):
        doc.add_paragraph("The AI payload was summarized/truncated to fit model request limits; Plotly charts use the full trace result.")

    doc.save(output_path)
    return output_path


def _add_hyperlink_paragraph(doc: Any, label: str, path: Path) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: {path.name}")


def _write_minimal_docx_report(
    output_path: Path,
    provider: str,
    focus_wallet: str,
    ai_text: str,
    chart_paths: Dict[str, Path],
    compact_payload: Dict[str, Any],
) -> Path:
    paragraphs = [
        "AI Transaction Analysis",
        f"Generated: {datetime.now(timezone.utc).isoformat()}",
        f"AI provider: {provider}",
        f"Specified wallet / focus: {focus_wallet or 'Not available'}",
        "Interactive Plotly Graphs",
        "Open these HTML files from the export folder to inspect the interactive charts.",
        f"Transaction volume and estimated net worth: {chart_paths['trend_chart'].name}",
        f"Money flow Sankey: {chart_paths['sankey_chart'].name}",
        "AI Analysis",
    ]
    paragraphs.extend(block.strip() for block in ai_text.split("\n\n") if block.strip())
    paragraphs.extend(
        [
            "Trace Scope",
            f"Transactions supplied to app: {compact_payload.get('transaction_count_supplied')}",
            f"Transactions included in AI payload: {compact_payload.get('transaction_count_in_ai_payload')}",
        ]
    )
    if compact_payload.get("payload_truncated"):
        paragraphs.append("The AI payload was summarized/truncated to fit model request limits; Plotly charts use the full trace result.")

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:body>"
        + "".join(_docx_paragraph_xml(text) for text in paragraphs)
        + (
            "<w:sectPr>"
            '<w:pgSz w:w="12240" w:h="15840"/>'
            '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>'
            "</w:sectPr>"
        )
        + "</w:body></w:document>"
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    relationships = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", relationships)
        archive.writestr("word/document.xml", document_xml)
    return output_path


def _docx_paragraph_xml(text: str) -> str:
    escaped = xml_escape(str(text)).replace("\n", "</w:t></w:r></w:p><w:p><w:r><w:t>")
    return f"<w:p><w:r><w:t>{escaped}</w:t></w:r></w:p>"


def _sum_amounts(items: List[Dict[str, Any]]) -> float:
    total = 0.0
    for item in items:
        try:
            total += float(item.get("amount", 0) or 0)
        except Exception:
            continue
    return total


def _sum_for_address(items: List[Dict[str, Any]], address_norm: str) -> float:
    if not address_norm:
        return 0.0
    total = 0.0
    for item in items:
        if _normalize_address(item.get("address")) == address_norm:
            try:
                total += float(item.get("amount", 0) or 0)
            except Exception:
                continue
    return total


def _items_for_address(items: List[Dict[str, Any]], address_norm: str) -> List[Dict[str, Any]]:
    if not address_norm:
        return []
    return [item for item in items if _normalize_address(item.get("address")) == address_norm]


def _normalize_address(address: Any) -> str:
    value = str(address or "").strip()
    if value.lower().startswith("0x"):
        return value.lower()
    return value


def _display_address(address: Any) -> str:
    value = str(address or "").strip() or "unknown"
    if len(value) <= 18:
        return value
    return f"{value[:10]}...{value[-6:]}"


def _parse_timestamp(value: Any) -> datetime:
    if not value:
        return datetime.now(timezone.utc)
    text = str(value).strip().replace("Z", "+00:00")
    try:
        parsed = datetime.fromisoformat(text)
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed
    except Exception:
        return datetime.now(timezone.utc)


def _round_float(value: float) -> float:
    if not math.isfinite(value):
        return 0.0
    return round(float(value), 10)
